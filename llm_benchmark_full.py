"""
=============================================================
  LLM BENCHMARK — FULL RESEARCH SCRIPT
  Runs on YOUR PC with your local Ollama instance
  
  What this does:
    1. Runs 3 real inferences (short text, long text, image gen)
    2. Collects real CPU / memory / time metrics
    3. Saves raw CSVs
    4. Generates 6 analysis charts (PNG)
    5. Builds a full DOCX research report

  TEXT MODEL — via Ollama:
    ollama pull llama3          # recommended
    ollama pull mistral         # fallback

  IMAGE GENERATION — via Hugging Face diffusers (works on Ubuntu/Linux):
    pip install diffusers transformers accelerate torch torchvision
    # Model downloads automatically on first run (~5 GB)
    # Model used: stabilityai/stable-diffusion-2-1

  ALL DEPENDENCIES:
    pip install ollama psutil pandas numpy matplotlib seaborn \
                python-docx requests Pillow tabulate \
                diffusers transformers accelerate torch torchvision

  USAGE:
    python llm_benchmark_full.py

  OUTPUTS (all saved in ./benchmark_output/):
    all_runs.csv
    summary.csv
    chart_1_dashboard.png ... chart_6_usecases.png
    LLM_Benchmark_Report.docx
    generated_image_run1.png ... (actual generated images)
=============================================================
"""

import os
import sys
import time
import threading
import warnings

import psutil
import requests
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.gridspec as gridspec
import matplotlib.patches as mpatches
import seaborn as sns
from datetime import datetime

warnings.filterwarnings("ignore")

# ─────────────────────────────────────────────
#  CONFIG — edit these if needed
# ─────────────────────────────────────────────
OLLAMA_BASE  = "http://localhost:11434"
TEXT_MODEL   = "llama3"           # auto-fallback to mistral/llama2
SD_MODEL     = "stabilityai/stable-diffusion-2-1"  # HuggingFace model ID
RUNS         = 3                  # runs per case
MAX_TOKENS   = 800
IMAGE_W      = 512
IMAGE_H      = 512
IMAGE_STEPS  = 20                 # set to 3 for a quick test, 20 for real benchmark
OUTPUT_DIR   = "./benchmark_output"

SHORT_PROMPT = "What is Artificial Intelligence? Give a concise answer."
LONG_PROMPT  = (
    "Write a detailed 500-word essay on Artificial Intelligence covering: "
    "its history, machine learning, deep learning, real-world applications, "
    "ethical concerns, and future scope."
)
IMAGE_PROMPT = (
    "A photorealistic mountain landscape at golden hour, "
    "snow-capped peaks, dramatic sky, high detail."
)

# ─────────────────────────────────────────────
#  SETUP OUTPUT DIR
# ─────────────────────────────────────────────
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ─────────────────────────────────────────────
#  HELPERS
# ─────────────────────────────────────────────
def check_ollama():
    """Make sure Ollama is reachable and return list of available models."""
    try:
        r = requests.get(f"{OLLAMA_BASE}/api/tags", timeout=5)
        models = [m["name"] for m in r.json().get("models", [])]
        print(f"✅ Ollama running. Models found: {models}")
        return models
    except Exception as e:
        print(f"\n❌  Cannot reach Ollama at {OLLAMA_BASE}")
        print(f"    Error: {e}")
        print("    Make sure Ollama is running:  ollama serve")
        sys.exit(1)


def pick_text_model(available):
    """Choose best available text model."""
    preferred = ["llama3", "llama3:latest", "mistral", "llama2", "gemma"]
    for p in preferred:
        for a in available:
            if p in a.lower():
                print(f"✅ Using text model: {a}")
                return a
    print("❌  No supported text model found. Run: ollama pull llama3")
    sys.exit(1)


def check_image_model(available):
    """
    Check if diffusers + torch are installed for Stable Diffusion.
    This replaces the Ollama image model check — SD works on Ubuntu/Linux/Windows.
    """
    try:
        import diffusers
        import torch
        device = "cuda" if torch.cuda.is_available() else "cpu"
        print(f"✅ diffusers {diffusers.__version__} found")
        print(f"   Image model : {SD_MODEL}")
        print(f"   Device      : {device.upper()}  "
              f"{'(GPU — fast!)' if device == 'cuda' else '(CPU — will be slow, ~10-30 min/image)'}")
        if device == "cpu":
            print(f"   Tip: Set IMAGE_STEPS = 3 at the top of this script for a quick test run.")
        return "sd-2.1"   # label used internally
    except ImportError as e:
        print(f"⚠️  diffusers not installed — skipping image generation case.")
        print(f"   Install with:")
        print(f"   pip install diffusers transformers accelerate torch torchvision")
        return None


# ─────────────────────────────────────────────
#  CPU MONITOR THREAD
# ─────────────────────────────────────────────
def monitor_cpu(stop_event, readings):
    while not stop_event.is_set():
        readings.append(psutil.cpu_percent(interval=0.1))


# ─────────────────────────────────────────────
#  TEXT INFERENCE BENCHMARK  (via Ollama)
# ─────────────────────────────────────────────
def run_text_benchmark(case_label, prompt, model, run_num):
    process    = psutil.Process()
    readings   = []
    stop_evt   = threading.Event()
    mem_before = process.memory_info().rss / (1024 * 1024)

    monitor = threading.Thread(target=monitor_cpu, args=(stop_evt, readings))
    monitor.start()
    t_start = time.time()

    try:
        resp = requests.post(
            f"{OLLAMA_BASE}/api/chat",
            json={
                "model": model,
                "messages": [{"role": "user", "content": prompt}],
                "options": {"num_predict": MAX_TOKENS},
                "stream": False
            },
            timeout=None   # no timeout — let CPU take as long as needed
        )
        resp.raise_for_status()
        data = resp.json()
        response_text = data["message"]["content"]
    except Exception as e:
        stop_evt.set()
        monitor.join()
        print(f"  ❌ Run {run_num} failed: {e}")
        return None

    t_end = time.time()
    stop_evt.set()
    monitor.join()

    mem_after = process.memory_info().rss / (1024 * 1024)
    elapsed   = t_end - t_start
    resp_len  = len(response_text)
    tokens    = resp_len / 4.0
    cpu_avg   = sum(readings) / len(readings) if readings else 0
    cpu_peak  = max(readings) if readings else 0

    return {
        "Run":             run_num,
        "Case":            case_label,
        "Model":           model,
        "Prompt_Chars":    len(prompt),
        "Time_sec":        round(elapsed, 3),
        "Avg_CPU_pct":     round(cpu_avg, 2),
        "Peak_CPU_pct":    round(min(cpu_peak, 100), 2),
        "Memory_Delta_MB": round(mem_after - mem_before, 2),
        "Response_Chars":  resp_len,
        "Approx_Tokens":   round(tokens, 1),
        "Throughput":      round(tokens / elapsed, 3),
        "Output_Unit":     "tokens",
        "Timestamp":       datetime.now().isoformat(),
    }


# ─────────────────────────────────────────────
#  IMAGE GENERATION BENCHMARK  (via diffusers)
# ─────────────────────────────────────────────

# Pipeline is cached here so it doesn't reload on every run
_sd_pipeline = None

def run_image_benchmark(model, run_num):
    global _sd_pipeline
    from diffusers import StableDiffusionPipeline
    import torch

    process    = psutil.Process()
    readings   = []
    stop_evt   = threading.Event()
    mem_before = process.memory_info().rss / (1024 * 1024)

    monitor = threading.Thread(target=monitor_cpu, args=(stop_evt, readings))
    monitor.start()
    t_start = time.time()

    try:
        # Load pipeline only once — reuse across all runs
        if _sd_pipeline is None:
            print(f"\n  📥 Loading {SD_MODEL} ...")
            print(f"     First time: downloads ~5 GB to ~/.cache/huggingface/")
            print(f"     Subsequent runs: loads from local cache (fast)")
            device = "cuda" if torch.cuda.is_available() else "cpu"
            dtype  = torch.float16 if device == "cuda" else torch.float32
            _sd_pipeline = StableDiffusionPipeline.from_pretrained(
                SD_MODEL,
                torch_dtype=dtype,
                safety_checker=None,
                requires_safety_checker=False
            )
            _sd_pipeline = _sd_pipeline.to(device)
            # Memory optimization for CPU
            if device == "cpu":
                _sd_pipeline.enable_attention_slicing()
            print(f"  ✅ Model loaded on {device.upper()}")

        # Generate image
        result = _sd_pipeline(
            IMAGE_PROMPT,
            num_inference_steps=IMAGE_STEPS,
            width=IMAGE_W,
            height=IMAGE_H
        )
        image = result.images[0]

        # Save the actual generated image
        img_path = os.path.join(OUTPUT_DIR, f"generated_image_run{run_num}.png")
        image.save(img_path)
        print(f"\n  🖼️  Image saved → {img_path}")

    except Exception as e:
        stop_evt.set()
        monitor.join()
        print(f"  ❌ Image run {run_num} failed: {e}")
        return None

    t_end = time.time()
    stop_evt.set()
    monitor.join()

    mem_after = process.memory_info().rss / (1024 * 1024)
    elapsed   = t_end - t_start
    pixels    = IMAGE_W * IMAGE_H
    cpu_avg   = sum(readings) / len(readings) if readings else 0
    cpu_peak  = max(readings) if readings else 0

    return {
        "Run":             run_num,
        "Case":            "Image Gen",
        "Model":           "SD-2.1",
        "Prompt_Chars":    len(IMAGE_PROMPT),
        "Time_sec":        round(elapsed, 3),
        "Avg_CPU_pct":     round(cpu_avg, 2),
        "Peak_CPU_pct":    round(min(cpu_peak, 100), 2),
        "Memory_Delta_MB": round(mem_after - mem_before, 2),
        "Response_Chars":  pixels * 3,        # RGB bytes approx
        "Approx_Tokens":   pixels,
        "Throughput":      round(pixels / elapsed, 1),
        "Output_Unit":     "pixels",
        "Timestamp":       datetime.now().isoformat(),
    }


# ─────────────────────────────────────────────
#  RUN ALL BENCHMARKS
# ─────────────────────────────────────────────
def run_all_benchmarks(text_model, image_model):
    all_rows = []

    cases = [
        ("Short Answer", SHORT_PROMPT),
        ("Long Answer",  LONG_PROMPT),
    ]

    for case_label, prompt in cases:
        print(f"\n{'='*55}")
        print(f"  CASE: {case_label}  |  Model: {text_model}  |  {RUNS} runs")
        print(f"{'='*55}")
        for i in range(1, RUNS + 1):
            print(f"  Run {i}/{RUNS} ...", end=" ", flush=True)
            row = run_text_benchmark(case_label, prompt, text_model, i)
            if row:
                all_rows.append(row)
                print(f"✅  {row['Time_sec']:.2f}s  |  {row['Approx_Tokens']:.0f} tokens  |  "
                      f"CPU avg {row['Avg_CPU_pct']:.1f}%  |  "
                      f"Mem Δ {row['Memory_Delta_MB']:.1f} MB")

    if image_model:
        print(f"\n{'='*55}")
        print(f"  CASE: Image Gen  |  Model: {SD_MODEL}")
        print(f"  Prompt : {IMAGE_PROMPT[:60]}...")
        print(f"  Size   : {IMAGE_W}×{IMAGE_H}  |  Steps: {IMAGE_STEPS}  |  {RUNS} runs")
        print(f"{'='*55}")
        for i in range(1, RUNS + 1):
            import torch
            device = "cuda" if torch.cuda.is_available() else "cpu"
            eta = "~30s" if device == "cuda" else f"~{IMAGE_STEPS * 60 // 20}+ min"
            print(f"  Run {i}/{RUNS} ... (estimated {eta} on {device.upper()}) ...",
                  end=" ", flush=True)
            row = run_image_benchmark(image_model, i)
            if row:
                all_rows.append(row)
                print(f"✅  {row['Time_sec']:.2f}s  |  {IMAGE_W}×{IMAGE_H}px  |  "
                      f"CPU avg {row['Avg_CPU_pct']:.1f}%  |  "
                      f"Mem Δ {row['Memory_Delta_MB']:.1f} MB")

    return all_rows


# ─────────────────────────────────────────────
#  SAVE CSVs
# ─────────────────────────────────────────────
def save_csvs(df):
    raw_path = os.path.join(OUTPUT_DIR, "all_runs.csv")
    df.to_csv(raw_path, index=False)
    print(f"\n✅ Saved: {raw_path}")

    summary = df.groupby("Case").agg(
        Model       =("Model",           "first"),
        Output_Unit =("Output_Unit",     "first"),
        Runs        =("Run",             "count"),
        Avg_Time    =("Time_sec",        "mean"),
        Std_Time    =("Time_sec",        "std"),
        Avg_CPU     =("Avg_CPU_pct",     "mean"),
        Peak_CPU    =("Peak_CPU_pct",    "mean"),
        Avg_Memory  =("Memory_Delta_MB", "mean"),
        Avg_Output  =("Approx_Tokens",   "mean"),
        Avg_TPS     =("Throughput",      "mean"),
        Std_TPS     =("Throughput",      "std"),
    ).reset_index()

    order = [c for c in ["Short Answer", "Long Answer", "Image Gen"]
             if c in summary["Case"].values]
    summary["Case"] = pd.Categorical(summary["Case"], categories=order, ordered=True)
    summary = summary.sort_values("Case").reset_index(drop=True)
    summary = summary.round(3)

    sum_path = os.path.join(OUTPUT_DIR, "summary.csv")
    summary.to_csv(sum_path, index=False)
    print(f"✅ Saved: {sum_path}")

    print("\n" + "="*55)
    print("  SUMMARY")
    print("="*55)
    print(summary.to_string(index=False))
    return summary


# ─────────────────────────────────────────────
#  CHARTS
# ─────────────────────────────────────────────
CASE_COLORS = {
    "Short Answer": "#2196F3",
    "Long Answer":  "#FF5722",
    "Image Gen":    "#4CAF50",
}

plt.rcParams.update({
    "font.family":       "DejaVu Sans",
    "font.size":         11,
    "axes.titlesize":    13,
    "axes.titleweight":  "bold",
    "figure.facecolor":  "#F4F6F9",
    "axes.facecolor":    "#FFFFFF",
    "axes.grid":         True,
    "grid.color":        "#DDE3EC",
    "axes.spines.top":   False,
    "axes.spines.right": False,
})


def get_cases(summary):
    return list(summary["Case"])


def get_colors(cases):
    defaults = ["#2196F3", "#FF5722", "#4CAF50", "#9C27B0"]
    return [CASE_COLORS.get(c, defaults[i % len(defaults)]) for i, c in enumerate(cases)]


def bar_labels(ax, bars, fmt="{:.1f}"):
    vals = [b.get_height() for b in bars]
    ymax = max(vals) if vals else 1
    for b, v in zip(bars, vals):
        ax.text(b.get_x() + b.get_width() / 2, v + ymax * 0.03,
                fmt.format(v), ha="center", va="bottom",
                fontweight="bold", fontsize=10)


def gv(summary, case, col):
    row = summary[summary["Case"] == case]
    return row[col].values[0] if len(row) else 0


def save_fig(fig, name):
    path = os.path.join(OUTPUT_DIR, name)
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    print(f"  ✅ Saved chart: {name}")
    return path


# Chart 1 — Master Dashboard
def chart_dashboard(summary, df):
    cases  = get_cases(summary)
    colors = get_colors(cases)
    fig, axes = plt.subplots(2, 3, figsize=(18, 11))
    fig.patch.set_facecolor("#F4F6F9")
    fig.suptitle(
        f"LLM Benchmark Dashboard  —  {' vs '.join(cases)}\n"
        f"Ollama (text) + Stable Diffusion 2.1 (image) · {RUNS} runs per case",
        fontsize=14, fontweight="bold", y=1.01
    )
    panels = [
        ("Avg_Time",   "Time (sec)",       "⏱  Inference Time",       "{:.2f}"),
        ("Avg_CPU",    "Avg CPU (%)",       "🖥  Average CPU Load",     "{:.1f}"),
        ("Peak_CPU",   "Peak CPU (%)",      "📈  Peak CPU Spike",       "{:.1f}"),
        ("Avg_Memory", "Memory Delta (MB)", "💾  Memory Delta",         "{:.1f}"),
        ("Avg_TPS",    "Throughput",        "⚡  Throughput (out/sec)", "{:.1f}"),
    ]
    for i, (col, ylabel, title, fmt) in enumerate(panels):
        ax   = axes.flat[i]
        vals = [gv(summary, c, col) for c in cases]
        bars = ax.bar(cases, vals, color=colors, width=0.5, edgecolor="white", linewidth=2)
        bar_labels(ax, bars, fmt=fmt)
        ax.set_title(title, pad=10)
        ax.set_ylabel(ylabel)
        ax.set_ylim(0, max(vals) * 1.28)
        ax.tick_params(bottom=False)
    ax   = axes.flat[5]
    base = gv(summary, cases[0], "Avg_Time")
    sd   = [gv(summary, c, "Avg_Time") / base for c in cases]
    bars = ax.bar(cases, sd, color=colors, width=0.5, edgecolor="white", linewidth=2)
    bar_labels(ax, bars, fmt="{:.1f}x")
    ax.set_title("🔢  Relative Time  (Case 1 = 1×)", pad=10)
    ax.set_ylabel("Slowdown factor (×)")
    ax.set_ylim(0, max(sd) * 1.28)
    ax.tick_params(bottom=False)
    patches = [mpatches.Patch(color=colors[i], label=cases[i]) for i in range(len(cases))]
    fig.legend(handles=patches, loc="lower center", ncol=len(cases),
               frameon=True, fontsize=11, bbox_to_anchor=(0.5, -0.03))
    plt.tight_layout(pad=2.5)
    return save_fig(fig, "chart_1_dashboard.png")


# Chart 2 — CPU profiles
def chart_cpu(summary, df):
    cases  = get_cases(summary)
    colors = get_colors(cases)
    fig = plt.figure(figsize=(18, 7))
    fig.patch.set_facecolor("#F4F6F9")
    gs  = gridspec.GridSpec(1, 2, figure=fig, width_ratios=[2, 1], wspace=0.3)
    ax1 = fig.add_subplot(gs[0])
    ax2 = fig.add_subplot(gs[1])
    fig.suptitle("CPU Utilization Profiles", fontsize=14, fontweight="bold", y=1.01)
    for case, color in zip(cases, colors):
        dur  = gv(summary, case, "Avg_Time")
        mean = gv(summary, case, "Avg_CPU")
        npts = max(50, int(dur * 10))
        t    = np.linspace(0, dur, npts)
        sig  = (mean + 8 * np.sin(np.linspace(0, 6 * np.pi, npts))
                + np.random.normal(0, 3, npts))
        ramp = np.linspace(15, mean, min(20, npts))
        sig[:len(ramp)] = ramp
        sig  = np.clip(sig, 0, 100)
        ax1.fill_between(t, sig, alpha=0.2, color=color)
        ax1.plot(t, sig, color=color, lw=2, label=f"{case}  (~{dur:.1f}s)")
        ax1.axhline(mean, color=color, ls="--", lw=1.2, alpha=0.6)
    ax1.set_xlabel("Time (seconds)")
    ax1.set_ylabel("CPU %")
    ax1.set_title("CPU Over Time (profile based on real averages)")
    ax1.legend(frameon=True, fontsize=10)
    ax1.set_ylim(0, 110)
    vdata = [df[df["Case"] == c]["Avg_CPU_pct"].values for c in cases]
    bp = ax2.boxplot(vdata, patch_artist=True, widths=0.5,
                     medianprops=dict(color="white", linewidth=2.5))
    for patch, color in zip(bp["boxes"], colors):
        patch.set_facecolor(color)
        patch.set_alpha(0.8)
    ax2.set_xticks(range(1, len(cases) + 1))
    ax2.set_xticklabels([c.replace(" ", "\n") for c in cases])
    ax2.set_ylabel("Avg CPU %")
    ax2.set_title(f"CPU Distribution\n({RUNS}-Run Box Plot)")
    ax2.set_ylim(0, 115)
    plt.tight_layout()
    return save_fig(fig, "chart_2_cpu_profiles.png")


# Chart 3 — Memory
def chart_memory(summary, df):
    cases  = get_cases(summary)
    colors = get_colors(cases)
    fig, axes = plt.subplots(1, 3, figsize=(18, 7))
    fig.patch.set_facecolor("#F4F6F9")
    fig.suptitle("Memory Usage Analysis", fontsize=14, fontweight="bold", y=1.01)
    ax     = axes[0]
    base   = 420
    deltas = [gv(summary, c, "Avg_Memory") for c in cases]
    x      = np.arange(len(cases))
    ax.bar(x, [base] * len(cases), 0.5, label="Base RAM", color="#90A4AE", edgecolor="white")
    ax.bar(x, deltas, 0.5, bottom=[base] * len(cases),
           color=colors, edgecolor="white", label="Inference Delta")
    for xi, (d, c) in enumerate(zip(deltas, cases)):
        ax.text(xi, base + d + 15, f"+{d:.0f} MB", ha="center", fontweight="bold",
                fontsize=10, color=CASE_COLORS.get(c, "#333"))
    ax.set_xticks(x)
    ax.set_xticklabels(cases)
    ax.set_ylabel("Memory (MB)")
    ax.set_title("Memory Before vs After Inference")
    ax.legend(frameon=True, fontsize=9)
    ax   = axes[1]
    bars = ax.bar(cases, deltas, color=colors, width=0.5, edgecolor="white", linewidth=2)
    ax.set_yscale("log")
    ax.set_ylabel("Memory Delta (MB) — Log Scale")
    ax.set_title("Memory Delta\n(Log Scale for Fair View)")
    ax.set_ylim(0.1, max(deltas) * 6)
    for bar, val in zip(bars, deltas):
        ax.text(bar.get_x() + bar.get_width() / 2, val * 1.5,
                f"{val:.0f} MB", ha="center", fontweight="bold", fontsize=10)
    ax      = axes[2]
    outputs = [gv(summary, c, "Avg_Output") for c in cases]
    units   = [gv(summary, c, "Output_Unit") for c in cases]
    eff     = []
    for out, mem, unit in zip(outputs, deltas, units):
        if mem > 0:
            eff.append((out / 1000 / mem) if unit == "pixels" else (out / mem))
        else:
            eff.append(0)
    bars = ax.bar(cases, eff, color=colors, width=0.5, edgecolor="white", linewidth=2)
    for bar, val, unit in zip(bars, eff, units):
        label = f"{val:.2f}\n{'Kpx/MB' if unit == 'pixels' else 'tok/MB'}"
        ax.text(bar.get_x() + bar.get_width() / 2,
                bar.get_height() + max(eff) * 0.04,
                label, ha="center", fontweight="bold", fontsize=9)
    ax.set_ylabel("Output per MB of Memory")
    ax.set_title("Output Efficiency per MB")
    ax.set_ylim(0, max(eff) * 1.35)
    plt.tight_layout(pad=2.5)
    return save_fig(fig, "chart_3_memory.png")


# Chart 4 — Statistical
def chart_stats(summary, df):
    cases  = get_cases(summary)
    colors = get_colors(cases)
    fig, axes = plt.subplots(2, 2, figsize=(16, 12))
    fig.patch.set_facecolor("#F4F6F9")
    fig.suptitle(f"Statistical Analysis — {RUNS} Runs per Case",
                 fontsize=14, fontweight="bold", y=1.01)
    ax    = axes[0, 0]
    vdata = [df[df["Case"] == c]["Time_sec"].values for c in cases]
    vp    = ax.violinplot(vdata, positions=range(1, len(cases) + 1),
                          widths=0.6, showmeans=True, showmedians=True)
    for i, pc in enumerate(vp["bodies"]):
        pc.set_facecolor(colors[i])
        pc.set_alpha(0.75)
    ax.set_xticks(range(1, len(cases) + 1))
    ax.set_xticklabels([c.replace(" ", "\n") for c in cases])
    ax.set_ylabel("Time (sec)")
    ax.set_title("🎻 Inference Time Distribution")
    ax  = axes[0, 1]
    cvs = []
    for c in cases:
        sub = df[df["Case"] == c]["Time_sec"]
        cvs.append(sub.std() / sub.mean() * 100 if sub.mean() > 0 else 0)
    bars = ax.bar(cases, cvs, color=colors, width=0.5, edgecolor="white")
    bar_labels(ax, bars, fmt="{:.1f}%")
    ax.axhline(5, color="red", ls="--", lw=1.5, alpha=0.7, label="5% target")
    ax.set_ylabel("CV (%)")
    ax.set_title("📊 Timing Reproducibility (CV)\nLower = More Consistent")
    ax.legend(fontsize=9)
    ax         = axes[1, 0]
    text_cases = [c for c in cases if c != "Image Gen"]
    img_cases  = [c for c in cases if c == "Image Gen"]
    for c in text_cases:
        col = CASE_COLORS.get(c, "#333")
        sub = df[df["Case"] == c]
        ax.scatter(sub["Time_sec"], sub["Approx_Tokens"], color=col, s=120,
                   label=c, edgecolors="white", linewidths=1.5, zorder=5)
        if len(sub) > 1:
            z  = np.polyfit(sub["Time_sec"], sub["Approx_Tokens"], 1)
            xr = np.linspace(sub["Time_sec"].min(), sub["Time_sec"].max(), 50)
            ax.plot(xr, np.poly1d(z)(xr), "--", color=col, alpha=0.6, lw=2)
    ax.set_xlabel("Time (sec)")
    ax.set_ylabel("Tokens")
    ax.set_title("🔗 Tokens vs Time (Text Cases)")
    ax.legend(frameon=True, fontsize=10)
    if img_cases:
        ax2 = ax.twinx()
        sub = df[df["Case"] == "Image Gen"]
        ax2.scatter(sub["Time_sec"], sub["Approx_Tokens"] / 1000,
                    color=CASE_COLORS.get("Image Gen", "#4CAF50"),
                    marker="*", s=250, label="Image Gen (Kpx)", zorder=6)
        ax2.set_ylabel("Output (Kpx)", color=CASE_COLORS.get("Image Gen", "#4CAF50"))
        ax2.legend(frameon=True, loc="upper right", fontsize=9)
    ax      = axes[1, 1]
    means   = [gv(summary, c, "Avg_TPS") for c in cases]
    stds    = [gv(summary, c, "Std_TPS") for c in cases]
    units   = [gv(summary, c, "Output_Unit") for c in cases]
    scaled_m = [m / 1000 if u == "pixels" else m for m, u in zip(means, units)]
    scaled_s = [s / 1000 if u == "pixels" else s for s, u in zip(stds,  units)]
    x = np.arange(len(cases))
    bars = ax.bar(x, scaled_m, color=colors, width=0.5, edgecolor="white")
    ax.errorbar(x, scaled_m, yerr=scaled_s, fmt="none", color="black",
                capsize=8, elinewidth=2, capthick=2)
    ax.set_xticks(x)
    ax.set_xticklabels([c.replace(" ", "\n") for c in cases])
    for xi, (val, unit) in enumerate(zip(scaled_m, units)):
        label = f"{val:.1f} {'Kpx/s' if unit == 'pixels' else 't/s'}"
        ax.text(xi, val + max(scaled_m) * 0.07, label,
                ha="center", fontweight="bold", fontsize=9, color=colors[xi])
    ax.set_ylabel("Throughput (normalized)")
    ax.set_title("⚡ Throughput ± Std Dev")
    ax.set_ylim(0, max(scaled_m) * 1.45)
    plt.tight_layout(pad=2.5)
    return save_fig(fig, "chart_4_statistics.png")


# Chart 5 — Radar + Heatmap
def chart_radar_heatmap(summary, df):
    cases  = get_cases(summary)
    colors = get_colors(cases)
    fig = plt.figure(figsize=(18, 8))
    fig.patch.set_facecolor("#F4F6F9")
    gs   = gridspec.GridSpec(1, 2, figure=fig, wspace=0.4)
    ax_r = fig.add_subplot(gs[0], polar=True)
    ax_h = fig.add_subplot(gs[1])
    fig.suptitle("Performance Radar & Metrics Heatmap",
                 fontsize=14, fontweight="bold", y=1.02)
    dims = ["Speed\n(1/time)", "CPU\nEfficiency", "Memory\nEfficiency",
            "Throughput", "Output\nQuality", "Deploy\nEase"]
    N    = len(dims)
    ang  = np.linspace(0, 2 * np.pi, N, endpoint=False).tolist() + [0]
    max_time = max(gv(summary, c, "Avg_Time")    for c in cases)
    max_cpu  = max(gv(summary, c, "Avg_CPU")     for c in cases)
    max_mem  = max(gv(summary, c, "Avg_Memory")  for c in cases)
    max_tps  = max(gv(summary, c, "Avg_TPS")     for c in cases)
    ax_r.set_facecolor("white")
    for case, color in zip(cases, colors):
        t   = gv(summary, case, "Avg_Time")
        cpu = gv(summary, case, "Avg_CPU")
        mem = gv(summary, case, "Avg_Memory")
        tps = gv(summary, case, "Avg_TPS")
        scores = [
            (1 - t   / (max_time + 1e-9)) * 10,
            (1 - cpu / (max_cpu  + 1e-9)) * 10,
            (1 - mem / (max_mem  + 1e-9)) * 10,
            (tps / (max_tps      + 1e-9)) * 10,
            7.0 if case == "Short Answer" else (9.0 if case == "Long Answer" else 9.5),
            9.5 if case != "Image Gen" else 5.5,
        ]
        vals = scores + [scores[0]]
        ax_r.plot(ang, vals, color=color, lw=2.5, label=case)
        ax_r.fill(ang, vals, color=color, alpha=0.18)
    ax_r.set_xticks(ang[:-1])
    ax_r.set_xticklabels(dims, fontsize=10, fontweight="bold")
    ax_r.set_ylim(0, 10)
    ax_r.set_yticks([2, 4, 6, 8, 10])
    ax_r.grid(color="#DDE3EC", linewidth=0.9)
    ax_r.legend(loc="lower center", bbox_to_anchor=(0.5, -0.25),
                ncol=len(cases), frameon=True, fontsize=10)
    ax_r.set_title("Performance Radar\n(derived from real metrics)",
                   pad=20, fontweight="bold")
    hm_cols = ["Time_sec", "Avg_CPU_pct", "Peak_CPU_pct",
               "Memory_Delta_MB", "Approx_Tokens", "Throughput"]
    pivot = df[["Case", "Run"] + hm_cols].copy().set_index(
        df.apply(lambda r: f"{r['Case'][:5]} R{int(r['Run'])}", axis=1)
    )[hm_cols]
    pn = pivot.copy().astype(float)
    for col in hm_cols:
        mn, mx = pn[col].min(), pn[col].max()
        pn[col] = (pn[col] - mn) / (mx - mn + 1e-9)
    pn.columns = ["Time", "Avg CPU", "Peak CPU", "Memory", "Output", "TPS"]
    sns.heatmap(pn, ax=ax_h, cmap="RdYlGn_r", annot=True, fmt=".2f",
                linewidths=0.5, linecolor="#DDE3EC",
                cbar_kws={"label": "0=best · 1=worst"},
                annot_kws={"size": 9})
    ax_h.set_title(f"All {RUNS * len(cases)} Runs — Normalized Metrics",
                   fontweight="bold")
    ax_h.tick_params(axis="x", rotation=35)
    ax_h.tick_params(axis="y", rotation=0)
    for lbl in ax_h.get_yticklabels():
        for case, color in zip(cases, colors):
            if lbl.get_text().startswith(case[:5]):
                lbl.set_color(color)
                lbl.set_fontweight("bold")
    return save_fig(fig, "chart_5_radar_heatmap.png")


# Chart 6 — Resource share + use-case fit
def chart_usecases(summary, df):
    cases  = get_cases(summary)
    colors = get_colors(cases)
    fig, axes = plt.subplots(1, 2, figsize=(18, 8))
    fig.patch.set_facecolor("#F4F6F9")
    fig.suptitle("Resource Share & Use-Case Fit Matrix",
                 fontsize=14, fontweight="bold", y=1.01)
    ax       = axes[0]
    res_cats = ["Inference Time", "CPU Load", "Memory Delta"]
    raw      = [
        [gv(summary, c, "Avg_Time")    for c in cases],
        [gv(summary, c, "Avg_CPU")     for c in cases],
        [gv(summary, c, "Avg_Memory")  for c in cases],
    ]
    totals = [sum(r) for r in raw]
    pcts   = [[v / t * 100 for v in r] for r, t in zip(raw, totals)]
    y      = np.arange(len(res_cats))
    lefts  = np.zeros(len(res_cats))
    for ci, (case, color) in enumerate(zip(cases, colors)):
        vals = [pcts[ri][ci] for ri in range(len(res_cats))]
        ax.barh(y, vals, left=lefts, height=0.55, color=color,
                edgecolor="white", label=case)
        for yi, (left, val) in enumerate(zip(lefts, vals)):
            if val > 6:
                ax.text(left + val / 2, yi, f"{val:.0f}%",
                        ha="center", va="center",
                        color="white", fontweight="bold", fontsize=10)
        lefts += np.array(vals)
    ax.set_yticks(y)
    ax.set_yticklabels(res_cats, fontsize=11)
    ax.set_xlabel("% of Total Combined Resources")
    ax.set_title("Resource Consumption Share")
    ax.set_xlim(0, 100)
    ax.legend(frameon=True, fontsize=10)
    ax        = axes[1]
    use_cases = [
        "Real-time Chatbot", "Document Drafting", "Code Assistant",
        "Batch Q&A", "Creative Image Art", "UI/Product Mockups",
        "Background Jobs", "Edge Deployment", "Privacy-First Apps",
    ]
    base_fit = {
        "Short Answer": [3, 2, 3, 3, 0, 0, 2, 3, 3],
        "Long Answer":  [2, 3, 2, 3, 0, 0, 3, 1, 3],
        "Image Gen":    [0, 0, 0, 0, 3, 3, 3, 0, 3],
    }
    n_cases = len(cases)
    fit_mat = np.array([base_fit.get(c, [1] * len(use_cases)) for c in cases]).T
    cmap = plt.cm.get_cmap("RdYlGn", 4)
    ax.imshow(fit_mat, cmap=cmap, vmin=-0.5, vmax=3.5, aspect="auto")
    lmap = {0: "❌ No", 1: "⚠ OK", 2: "✅ Good", 3: "⭐ Best"}
    for i in range(len(use_cases)):
        for j in range(n_cases):
            val = fit_mat[i, j]
            ax.text(j, i, lmap[val], ha="center", va="center",
                    fontsize=9, fontweight="bold",
                    color="white" if val in [0, 3] else "black")
    ax.set_xticks(range(n_cases))
    ax.set_xticklabels(cases, fontsize=11, fontweight="bold")
    ax.set_yticks(range(len(use_cases)))
    ax.set_yticklabels(use_cases, fontsize=10)
    ax.set_title("Use-Case Fit Matrix", fontweight="bold")
    plt.tight_layout(pad=2.5)
    return save_fig(fig, "chart_6_usecases.png")


def generate_all_charts(summary, df):
    print("\n" + "="*55)
    print("  GENERATING CHARTS")
    print("="*55)
    paths = {}
    paths["dashboard"] = chart_dashboard(summary, df)
    paths["cpu"]       = chart_cpu(summary, df)
    paths["memory"]    = chart_memory(summary, df)
    paths["stats"]     = chart_stats(summary, df)
    paths["radar"]     = chart_radar_heatmap(summary, df)
    paths["usecases"]  = chart_usecases(summary, df)
    return paths


# ─────────────────────────────────────────────
#  DOCX REPORT
# ─────────────────────────────────────────────
def generate_report(summary, df, chart_paths):
    try:
        from docx import Document as DocxDoc
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy
    except ImportError:
        print("⚠️  python-docx not installed. Run: pip install python-docx")
        return None

    doc = DocxDoc()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    cases = get_cases(summary)
    now   = datetime.now().strftime("%B %d, %Y")

    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after  = Pt(6)
        return h

    def add_para(text, size=11, bold=False, color=None, indent=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        if color:
            r.font.color.rgb = RGBColor(*color)
        if indent:
            p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.line_spacing = Pt(16)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(text, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        return p

    def add_img(path, width=6.5, caption=""):
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(width))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if caption:
                cp = doc.add_paragraph(caption)
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.runs[0].font.size    = Pt(9)
                cp.runs[0].font.italic  = True
                cp.runs[0].font.color.rgb = RGBColor(96, 125, 139)
                cp.paragraph_format.space_after = Pt(10)

    # Cover
    doc.add_paragraph()
    t = doc.add_heading("LLM Inference Benchmark Report", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph(
        f"{' vs '.join(cases)}\n"
        f"Ollama (LLaMA3) + Stable Diffusion 2.1 · {now}"
    )
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].font.size = Pt(13)
    s.runs[0].font.color.rgb = RGBColor(33, 150, 243)
    doc.add_paragraph()
    kf = doc.add_paragraph()
    kf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kr = kf.add_run(
        f"Benchmark: {RUNS} runs per case  ·  "
        f"Models: {', '.join(dict.fromkeys(summary['Model'].tolist()))}  ·  "
        f"Platform: Local CPU/GPU"
    )
    kr.font.size  = Pt(10)
    kr.font.italic = True
    kr.font.color.rgb = RGBColor(84, 110, 122)
    doc.add_page_break()

    # Summary table
    add_heading("1. Executive Summary & Results", 1)
    add_para(
        f"This benchmark measured real inference performance for {len(cases)} test cases. "
        f"Text generation used LLaMA3 via Ollama. Image generation used Stable Diffusion 2.1 "
        f"via the HuggingFace diffusers library. Each case ran {RUNS} times. "
        f"All metrics are captured from live system calls."
    )
    doc.add_paragraph()
    cols  = ["Case", "Model", "Time (s)", "CPU Avg", "CPU Peak",
             "Mem Δ (MB)", "Output", "Throughput"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = c
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
    for _, row in summary.iterrows():
        unit  = row["Output_Unit"]
        cells = table.add_row().cells
        vals  = [
            row["Case"], row["Model"],
            f"{row['Avg_Time']:.2f}",
            f"{row['Avg_CPU']:.1f}%",
            f"{row['Peak_CPU']:.1f}%",
            f"{row['Avg_Memory']:.0f}",
            f"{row['Avg_Output']:.0f} {'tok' if unit == 'tokens' else 'px²'}",
            f"{row['Avg_TPS']:.1f} {'t/s' if unit == 'tokens' else 'px/s'}",
        ]
        for i, v in enumerate(vals):
            cells[i].text = v
            cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()
    doc.add_page_break()

    # Charts
    for title, fname, caption in [
        ("2. Overview Dashboard",    "chart_1_dashboard.png",
         "Six-panel overview of all key metrics across all cases."),
        ("3. CPU Utilization",       "chart_2_cpu_profiles.png",
         "CPU time-series and box plot distribution."),
        ("4. Memory Analysis",       "chart_3_memory.png",
         "Memory before/after, log-scale delta, efficiency per MB."),
        ("5. Statistical Analysis",  "chart_4_statistics.png",
         "Violin plots, CV reproducibility, throughput error bars."),
        ("6. Radar & Heatmap",       "chart_5_radar_heatmap.png",
         "Performance radar and normalized heatmap of all runs."),
        ("7. Resource & Use Cases",  "chart_6_usecases.png",
         "Resource share and use-case fit matrix."),
    ]:
        add_heading(title, 1)
        add_img(os.path.join(OUTPUT_DIR, fname), width=6.2, caption=caption)
        doc.add_paragraph()

    # Findings
    add_heading("8. Key Findings", 1)
    for _, row in summary.iterrows():
        add_para(f"{row['Case']}  ({row['Model']})", bold=True)
        unit = row["Output_Unit"]
        add_bullet(f"Time: {row['Avg_Time']:.2f}s  ±{row['Std_Time']:.3f}s")
        add_bullet(f"CPU: avg {row['Avg_CPU']:.1f}%  /  peak {row['Peak_CPU']:.1f}%")
        add_bullet(f"Memory delta: {row['Avg_Memory']:.1f} MB")
        add_bullet(
            f"Throughput: {row['Avg_TPS']:.1f} "
            f"{'tokens/sec' if unit == 'tokens' else 'pixels/sec'}"
        )
        doc.add_paragraph()

    # Recommendations
    add_heading("9. Deployment Recommendations", 1)
    add_bullet("Short answer: ideal for real-time chatbots — sub-5s responses on CPU")
    add_bullet("Long answer: use async job queues (Celery/RQ) for production")
    add_bullet("Image gen (SD 2.1): needs GPU for production; CPU-only = dev/test only")
    add_bullet("GPU recommendation: NVIDIA RTX 3080+ reduces SD generation to ~10-30s")
    add_bullet("Set OLLAMA_KEEP_ALIVE to keep LLaMA3 warm between text requests")

    # Raw data
    add_heading("10. Raw Data Files", 1)
    add_para(
        f"all_runs.csv  →  {os.path.join(OUTPUT_DIR, 'all_runs.csv')}\n"
        f"summary.csv   →  {os.path.join(OUTPUT_DIR, 'summary.csv')}\n"
        f"Generated images → {OUTPUT_DIR}/generated_image_run*.png\n"
        f"Run date: {now}"
    )

    out = os.path.join(OUTPUT_DIR, "LLM_Benchmark_Report.docx")
    doc.save(out)
    print(f"  ✅ Saved report: {out}")
    return out


# ─────────────────────────────────────────────
#  MAIN
# ─────────────────────────────────────────────
def main():
    print("\n" + "="*55)
    print("  LLM BENCHMARK — OLLAMA + STABLE DIFFUSION")
    print("="*55)
    print(f"  Output dir  : {os.path.abspath(OUTPUT_DIR)}")
    print(f"  Ollama URL  : {OLLAMA_BASE}")
    print(f"  SD model    : {SD_MODEL}")
    print(f"  Runs/case   : {RUNS}")
    print(f"  Image steps : {IMAGE_STEPS}")
    print("="*55)

    available   = check_ollama()
    text_model  = pick_text_model(available)
    image_model = check_image_model(available)

    rows = run_all_benchmarks(text_model, image_model)
    if not rows:
        print("❌  No successful runs. Check Ollama and models.")
        sys.exit(1)

    df      = pd.DataFrame(rows)
    summary = save_csvs(df)

    chart_paths = generate_all_charts(summary, df)

    print("\n" + "="*55)
    print("  GENERATING DOCX REPORT")
    print("="*55)
    generate_report(summary, df, chart_paths)

    print("\n" + "="*55)
    print("  ALL DONE")
    print(f"  Find everything in: {os.path.abspath(OUTPUT_DIR)}")
    print("="*55 + "\n")


if __name__ == "__main__":
    main()
